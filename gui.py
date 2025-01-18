# Python Imports
from pathlib import Path
from tkinter import Tk,ttk,font ,Canvas, Entry, Text, Button, PhotoImage,OptionMenu,StringVar,END,Frame,messagebox,filedialog
import tkinter as tk
import numpy as np
import pandas as pd
from docx import Document
from textwrap import wrap
import csv
from docx2pdf import convert

OUTPUT_PATH = Path(__file__).parent
ASSETS_PATH = OUTPUT_PATH / Path(r"C:\Users\M Sami Furqan\Desktop\GUI\build\assets\frame0")


def relative_to_assets(path: str) -> Path:
    return ASSETS_PATH / Path(path)

#defining the list of symptoms
#----------------------------------------------------------------------------------------------------------------------------------------#
l1=['back_pain','constipation','abdominal_pain','diarrhoea','mild_fever','yellow_urine',
'yellowing_of_eyes','acute_liver_failure','fluid_overload','swelling_of_stomach',
'swelled_lymph_nodes','malaise','blurred_and_distorted_vision','phlegm','throat_irritation',
'redness_of_eyes','sinus_pressure','runny_nose','congestion','chest_pain','weakness_in_limbs',
'fast_heart_rate','pain_during_bowel_movements','pain_in_anal_region','bloody_stool',
'irritation_in_anus','neck_pain','dizziness','cramps','bruising','obesity','swollen_legs',
'swollen_blood_vessels','puffy_face_and_eyes','enlarged_thyroid','brittle_nails',
'swollen_extremeties','excessive_hunger','extra_marital_contacts','drying_and_tingling_lips',
'slurred_speech','knee_pain','hip_joint_pain','muscle_weakness','stiff_neck','swelling_joints',
'movement_stiffness','spinning_movements','loss_of_balance','unsteadiness',
'weakness_of_one_body_side','loss_of_smell','bladder_discomfort','foul_smell_of urine',
'continuous_feel_of_urine','passage_of_gases','internal_itching','toxic_look_(typhos)',
'depression','irritability','muscle_pain','altered_sensorium','red_spots_over_body','belly_pain',
'abnormal_menstruation','dischromic _patches','watering_from_eyes','increased_appetite','polyuria','family_history','mucoid_sputum',
'rusty_sputum','lack_of_concentration','visual_disturbances','receiving_blood_transfusion',
'receiving_unsterile_injections','coma','stomach_bleeding','distention_of_abdomen',
'history_of_alcohol_consumption','fluid_overload','blood_in_sputum','prominent_veins_on_calf',
'palpitations','painful_walking','pus_filled_pimples','blackheads','scurring','skin_peeling',
'silver_like_dusting','small_dents_in_nails','inflammatory_nails','blister','red_sore_around_nose',
'yellow_crust_ooze']
#----------------------------------------------------------------------------------------------------------------------------------------#

#defining the list of diseases
#----------------------------------------------------------------------------------------------------------------------------------------#
disease=['Fungal infection','Allergy','GERD','Chronic cholestasis','Drug Reaction',
'Peptic ulcer diseae','AIDS','Diabetes','Gastroenteritis','Bronchial Asthma','Hypertension',
' Migraine','Cervical spondylosis',
'Paralysis (brain hemorrhage)','Jaundice','Malaria','Chicken pox','Dengue','Typhoid','hepatitis A',
'Hepatitis B','Hepatitis C','Hepatitis D','Hepatitis E','Alcoholic hepatitis','Tuberculosis',
'Common Cold','Pneumonia','Dimorphic hemmorhoids(piles)',
'Heartattack','Varicoseveins','Hypothyroidism','Hyperthyroidism','Hypoglycemia','Osteoarthristis',
'Arthritis','(vertigo) Paroymsal  Positional Vertigo','Acne','Urinary tract infection','Psoriasis',
'Impetigo']
#----------------------------------------------------------------------------------------------------------------------------------------#

l2=[]
for x in range(0,len(l1)):
    l2.append(0)

# TESTING DATA df -------------------------------------------------------------------------------------
df=pd.read_csv("Data\\Training.csv")

df.replace({'prognosis':{'Fungal infection':0,'Allergy':1,'GERD':2,'Chronic cholestasis':3,'Drug Reaction':4,
'Peptic ulcer diseae':5,'AIDS':6,'Diabetes ':7,'Gastroenteritis':8,'Bronchial Asthma':9,'Hypertension ':10,
'Migraine':11,'Cervical spondylosis':12,
'Paralysis (brain hemorrhage)':13,'Jaundice':14,'Malaria':15,'Chicken pox':16,'Dengue':17,'Typhoid':18,'hepatitis A':19,
'Hepatitis B':20,'Hepatitis C':21,'Hepatitis D':22,'Hepatitis E':23,'Alcoholic hepatitis':24,'Tuberculosis':25,
'Common Cold':26,'Pneumonia':27,'Dimorphic hemmorhoids(piles)':28,'Heart attack':29,'Varicose veins':30,'Hypothyroidism':31,
'Hyperthyroidism':32,'Hypoglycemia':33,'Osteoarthristis':34,'Arthritis':35,
'(vertigo) Paroymsal  Positional Vertigo':36,'Acne':37,'Urinary tract infection':38,'Psoriasis':39,
'Impetigo':40}},inplace=True)

# print(df.head())

X= df[l1]

y = df[["prognosis"]]
np.ravel(y)
# print(y)

# TRAINING DATA tr --------------------------------------------------------------------------------
tr=pd.read_csv("Data\\Testing.csv")
tr.replace({'prognosis':{'Fungal infection':0,'Allergy':1,'GERD':2,'Chronic cholestasis':3,'Drug Reaction':4,
'Peptic ulcer diseae':5,'AIDS':6,'Diabetes ':7,'Gastroenteritis':8,'Bronchial Asthma':9,'Hypertension ':10,
'Migraine':11,'Cervical spondylosis':12,
'Paralysis (brain hemorrhage)':13,'Jaundice':14,'Malaria':15,'Chicken pox':16,'Dengue':17,'Typhoid':18,'hepatitis A':19,
'Hepatitis B':20,'Hepatitis C':21,'Hepatitis D':22,'Hepatitis E':23,'Alcoholic hepatitis':24,'Tuberculosis':25,
'Common Cold':26,'Pneumonia':27,'Dimorphic hemmorhoids(piles)':28,'Heart attack':29,'Varicose veins':30,'Hypothyroidism':31,
'Hyperthyroidism':32,'Hypoglycemia':33,'Osteoarthristis':34,'Arthritis':35,
'(vertigo) Paroymsal  Positional Vertigo':36,'Acne':37,'Urinary tract infection':38,'Psoriasis':39,
'Impetigo':40}},inplace=True)

X_test= tr[l1] #test data features
y_test = tr[["prognosis"]] #test data labels
np.ravel(y_test) #raveling
# ------------------------------------------------------------------------------------------------------


window = Tk()

window.geometry("1117x550") # Set the window size
window.configure(bg = "#303030") # Set the background color
window.title("Disease Prediction System") # Set the window title
window.iconbitmap("assets\\Application Icon\\image-removebg-preview.ico")

# Define a custom font "Lexend".
custom_font = font.Font(family="Lexend", size=10, weight="bold")

# Setting up the global font
window.option_add("*Font", custom_font)

# Function to create a combobox at a specific position
def create_combobox(x1, y1, x2, y2,symptom_var, styleCB):
    # Calculate the center position for the combobox
    x = (x1 + x2) / 2
    y = (y1 + y2) / 2
    style = ttk.Style()
    # style.theme_use('default')
    symptom_entries = []
    style.configure("TCombobox",
                foreground="blue",  # Text color
                background="lightgray",  # Background color
                fieldbackground="white",  # Background inside the entry field
                borderwidth=3,
                )
    # Create a combobox with values as the position on the screen
    combobox = ttk.Combobox(window, values=l1, textvariable=symptom_var ,style=styleCB)
    combobox.place(x=x, y=y, anchor='center')
    # Place the combobox at the calculated position

# String Variables to store the selected symptoms.
Symptom1 = StringVar()
Symptom2 = StringVar()
Symptom3 = StringVar()
Symptom4 = StringVar()
Symptom5 = StringVar()
Name = StringVar()



def DecisionTree():
    """
    This function takes the selected symptoms from the user and uses them to predict a disease
    using a decision tree classifier. If no symptoms are selected, it shows an error message
    and returns. It calculates the accuracy of the model and prints it. It then predicts the
    disease based on the selected symptoms and updates the text of the canvas item accordingly.

    Parameters:
    None

    Returns:
    None
    """
    if [Symptom1.get(),Symptom2.get(),Symptom3.get(),Symptom4.get(),Symptom5.get()] == ['', '', '', '', '']: # Check if any symptoms are selected
        messagebox.showerror("Error", "Please select at least one symptom.")
        return
    from sklearn import tree
    from sklearn.metrics import accuracy_score
    
    clf3 = tree.DecisionTreeClassifier()  # empty model of the decision tree
    clf3 = clf3.fit(X, y)

    # calculating accuracy
    y_pred = clf3.predict(X_test)
    print(accuracy_score(y_test, y_pred))
    print(accuracy_score(y_test, y_pred, normalize=False))

    psymptoms = [Symptom1.get(), Symptom2.get(), Symptom3.get(), Symptom4.get(), Symptom5.get()]
    
    l2 = [0] * len(l1)
    for k in range(len(l1)):
        if l1[k] in psymptoms:
            l2[k] = 1

    inputtest = [l2]
    predict = clf3.predict(inputtest)
    predicted = predict[0]
    
    if predicted in range(len(disease)):
        canvas.itemconfig(tid_1, text=disease[predicted])  # Update the text of the canvas item
    else:
        canvas.itemconfig(tid_1, text="Not Found")


def randomforest():

    """
    This function uses a Random Forest classifier to predict a disease based on selected symptoms.
    It first checks if symptoms are selected, then trains the classifier with the given dataset.
    The function calculates and prints the accuracy of the model. It predicts the disease using
    the selected symptoms and updates the text of a canvas item with the predicted disease name.

    Parameters:
    None

    Returns:
    None
    """
    if [Symptom1.get(),Symptom2.get(),Symptom3.get(),Symptom4.get(),Symptom5.get()] == ['', '', '', '', '']: # Check if any symptoms are selected
        messagebox.showerror("Error", "Please select at least one symptom.")
        return
    
    from sklearn.ensemble import RandomForestClassifier
    clf4 = RandomForestClassifier()
    clf4 = clf4.fit(X,np.ravel(y))

    # calculating accuracy-------------------------------------------------------------------
    from sklearn.metrics import accuracy_score
    y_pred=clf4.predict(X_test)
    print(accuracy_score(y_test, y_pred))
    print(accuracy_score(y_test, y_pred,normalize=False))
    # ---------------------------------------------------------------------------------------

    psymptoms = [Symptom1.get(),Symptom2.get(),Symptom3.get(),Symptom4.get(),Symptom5.get()]

    for k in range(0,len(l1)):
        for z in psymptoms:
            if(z==l1[k]):
                l2[k]=1

    inputtest = [l2]
    predict = clf4.predict(inputtest)
    predicted=predict[0]

    h='no'
    for a in range(0,len(disease)):
        if(predicted == a):
            h='yes'
            break

    if (h=='yes'):
       canvas.itemconfig(tid_2, text=disease[a])
    else:
       pass



def NaiveBayes():
    
    """
    This function uses a Naive Bayes classifier to predict a disease based on selected symptoms.
    It first checks if symptoms are selected, then trains the classifier with the given dataset.
    The function calculates and prints the accuracy of the model. It predicts the disease using
    the selected symptoms and updates the text of a canvas item with the predicted disease name.

    Parameters:
    None

    Returns:
    None
    """
    if [Symptom1.get(),Symptom2.get(),Symptom3.get(),Symptom4.get(),Symptom5.get()] == ['', '', '', '', '']: # Check if any symptoms are selected
        messagebox.showerror("Error", "Please select at least one symptom.")
        return    
    
    from sklearn.naive_bayes import GaussianNB
    gnb = GaussianNB()
    gnb=gnb.fit(X,np.ravel(y))

    # calculating accuracy-------------------------------------------------------------------
    from sklearn.metrics import accuracy_score
    y_pred=gnb.predict(X_test)
    print(accuracy_score(y_test, y_pred))
    print(accuracy_score(y_test, y_pred,normalize=False))
    # ---------------------------------------------------------------------------------------

    psymptoms = [Symptom1.get(),Symptom2.get(),Symptom3.get(),Symptom4.get(),Symptom5.get()]
    for k in range(0,len(l1)):
        for z in psymptoms:
            if(z==l1[k]):
                l2[k]=1

    inputtest = [l2]
    predict = gnb.predict(inputtest)
    predicted=predict[0]

    h='no'
    for a in range(0,len(disease)):
        if(predicted == a):
            h='yes'
            break

    if (h=='yes'):
        canvas.itemconfig(tid_3, text=disease[a])
    
    else:
        pass

"""    
Rest of the UI Elements, Designed in Figma, Later converted into assets using Figma to Python Converter,
This include, Images, Lines, Texts, Fonts, Buttons.
More over, The certain functions are added later to the buttons, the code of which is also below.
""" 
canvas = Canvas(
    window,
    bg = "#303030",
    height = 550,
    width = 1117,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
)

canvas.place(x = 0, y = 0)
image_image_1 = PhotoImage(
    file=relative_to_assets("image_1.png"))
image_1 = canvas.create_image(
    536.0,
    317.0,
    image=image_image_1
)

image_image_2 = PhotoImage(
    file=relative_to_assets("image_2.png"))
image_2 = canvas.create_image(
    536.0,
    317.0,
    image=image_image_2
)

image_image_3 = PhotoImage(
    file=relative_to_assets("image_3.png"))
image_3 = canvas.create_image(
    175.0,
    231.0,
    image=image_image_3
)

image_image_4 = PhotoImage(
    file=relative_to_assets("image_4.png"))
image_4 = canvas.create_image(
    175.0,
    463.0,
    image=image_image_4
)

image_image_5 = PhotoImage(
    file=relative_to_assets("image_5.png"))
image_5 = canvas.create_image(
    919.0,
    197.0,
    image=image_image_5
)

image_image_6 = PhotoImage(
    file=relative_to_assets("image_6.png"))
image_6 = canvas.create_image(
    919.0,
    428.0,
    image=image_image_6
)

image_image_7 = PhotoImage(
    file=relative_to_assets("image_7.png"))
image_7 = canvas.create_image(
    558.0,
    41.0,
    image=image_image_7
)

canvas.create_text(
    148.0,
    26.0,
    anchor="nw",
    text="Disease Prediction from Symptoms - Machine Learning - Python",
    fill="#FFFFFF",
    font=("Lexend Bold", 25 * -1)
)

canvas.create_text(
    408.0,
    119.0,
    anchor="nw",
    text="Symptoms Selector",
    fill="#FFFFFF",
    font=("Lexend Bold", 27 * -1)
)


canvas.create_text(
    76.0,
    119.0,
    anchor="nw",
    text="Patient’s Info",
    fill="#FFFFFF",
    font=("Lexend Bold", 27 * -1)
)

canvas.create_text(
    54.0,
    404.0,
    anchor="nw",
    text="Generate Report",
    fill="#FFFFFF",
    font=("Lexend Bold", 27 * -1)
)

canvas.create_text(
    791.0,
    119.0,
    anchor="nw",
    text="Algorithm Selection",
    fill="#FFFFFF",
    font=("Lexend Bold", 27 * -1)
)

canvas.create_text(
    869.0,
    332.0,
    anchor="nw",
    text="Results",
    fill="#FFFFFF",
    font=("Lexend Bold", 27 * -1)
)

canvas.create_text(
    371.0,
    203.0,
    anchor="nw",
    text="Symptom 1 :",
    fill="#FFFFFF",
    font=("Lexend Regular", 20 * -1)
)

canvas.create_text(
    370.0,
    265.0,
    anchor="nw",
    text="Symptom 2 :",
    fill="#FFFFFF",
    font=("Lexend Regular", 20 * -1)
)

canvas.create_text(
    370.0,
    328.0,
    anchor="nw",
    text="Symptom 3 :",
    fill="#FFFFFF",
    font=("Lexend Regular", 20 * -1)
)

canvas.create_text(
    369.0,
    390.0,
    anchor="nw",
    text="Symptom 4 :",
    fill="#FFFFFF",
    font=("Lexend Regular", 20 * -1)
)

canvas.create_text(
    370.0,
    453.0,
    anchor="nw",
    text="Symptom 5 :",
    fill="#FFFFFF",
    font=("Lexend Regular", 20 * -1)
)


entry_image_1 = PhotoImage(
    file=relative_to_assets("entry_1.png"))
entry_bg_1 = canvas.create_image(
    222.5,
    217.5,
    image=entry_image_1
)
entry_1 = Text(
    bd=0,
    bg="#D9D9D9",
    fg="#000716",
    highlightthickness=0
)
def get_text():
    """
    Gets the text from the two text boxes and checks if they are empty and if the second one is a valid contact number.
    If both conditions are met, a message box is shown with the entered information. If not, an error message is shown.
    """
    global entrtxt1,entrtxt2
    entrtxt1 = entry_1.get("1.0", "end-1c")
    entrtxt2 = entry_2.get("1.0", "end-1c")
    if (entrtxt1 == "") or (entrtxt2 == ""):
        tk.messagebox.showerror("Error", "Please Enter Patient Name and Contact Number")
    if entrtxt2 is not None and not entrtxt2.isdigit():
        tk.messagebox.showerror("Error", "Please Enter Valid Contact Number")
    else:
        tk.messagebox.showinfo("Information Saved ", f"Patient Name : {entrtxt1}\nPateint's Contact : {entrtxt2}")
    
entry_1.place(
    x=141.0,
    y=207.0,
    width=163.0,
    height=25
)

button_image_1 = PhotoImage(
    file=relative_to_assets("button_1.png"))
button_1 = Button(
    image=button_image_1,
    borderwidth=0,
    highlightthickness=0,
    command=DecisionTree,
    relief="flat"
)
button_1.place(
    x=795.0,
    y=166.0,
    width=260.0,
    height=33.0
)

button_image_2 = PhotoImage(
    file=relative_to_assets("button_2.png"))
button_2 = Button(
    image=button_image_2,
    borderwidth=0,
    highlightthickness=0,
    command=randomforest,
    relief="flat"
)
button_2.place(
    x=795.0,
    y=208.0,
    width=260.0,
    height=33.0
)

button_image_3 = PhotoImage(
    file=relative_to_assets("button_3.png"))
button_3 = Button(
    image=button_image_3,
    borderwidth=0,
    highlightthickness=0,
    command=NaiveBayes,
    relief="flat"
)
button_3.place(
    x=795.0,
    y=250.0,
    width=260.0,
    height=33.0
)

button_image_4 = PhotoImage(
    file=relative_to_assets("button_4.png"))
button_4 = Button(
    image=button_image_4,
    borderwidth=0,
    highlightthickness=0,
    command=get_text,
    relief="flat"
)
button_4.place(
    x=46.0,
    y=315.0,
    width=260.0,
    height=33.0
)


def generate_report():
    """
    Generate a report based on the user's input and disease predictions from the algorithms. 
    The report includes the patient's name, contact, symptoms, and the diseases predicted by each 
    algorithm. Additionally, it provides recommendations for each disease predicted. 
    The report is saved as a DOCX file and the user can optionally convert it to PDF.

    Parameters:
    None

    Returns:
    None
    """
    if canvas.itemcget(tid_1, "text") == "." and canvas.itemcget(tid_2, "text") == "." and canvas.itemcget(tid_3, "text") == ".":
        messagebox.showerror("Error", "No Disease Detected, Please generate atleast one disease using Algorithm Selection") 
    else:
        pass
    from datetime import datetime
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
    from docx.shared import RGBColor, Pt
    global entrtxt1,entrtxt2
    # Create a new Word document
    document = Document()

    # Modify the Title style
    style = document.styles['Title']
    font = style.font
    font.name = "Century Gothic"
    font.size = Pt(24)
    font.color.rgb = RGBColor(0, 0, 0)
    font.bold = True

    # Adding and Styling the Title - Title of the Report
    H1 = document.add_heading(level=0)
    H1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = H1.add_run(f"Disease Prediction from Symptoms")
    run1.font.name = "Century Gothic"
    run1.font.size = Pt(24)
    run1.font.color.rgb = RGBColor(0, 0, 0)
    run1.font.bold = True
  
    # Adding and Styling the Subtitle - Report Title
    H2 = document.add_heading(level=1)
    H2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = H2.add_run(f"Report")
    run5.font.name = "Century Gothic"
    run5.font.size = Pt(20)
    run5.font.color.rgb = RGBColor(255, 0, 0)
    run5.font.bold = True
    run5.font.underline = True
    
    # Date time in the report 
    p1 = document.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    current_date = datetime.now()
    formatted_date = current_date.strftime("%m/%d/%Y")
    run4 = p1.add_run(f"Date : {formatted_date}")
    run4.font.name = "Century Gothic"
    run4.font.size = Pt(12)
    run4.font.color.rgb = RGBColor(0, 0, 0)
    run4.font.bold = True
    
    # Adding and Styling the Subtitle - Report Title
    H3 = document.add_heading(level=1)
    H3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run6 = H3.add_run(f"Pateints Details")
    run6.font.name = "Century Gothic"
    run6.font.size = Pt(20)
    run6.font.color.rgb = RGBColor(79, 129, 189)
    run6.font.bold = True
    
    # Add patient details with direct formatting
    paragraph1 = document.add_paragraph()
    run2 = paragraph1.add_run(f"Patient Name : {entrtxt1}")
    run2.font.name = "Century Gothic"
    run2.font.size = Pt(12)
    run2.font.bold = True

    # Add patient contact
    paragraph2 = document.add_paragraph()
    run3 = paragraph2.add_run(f"Patient's Contact : {entrtxt2}")
    run3.font.name = "Century Gothic"
    run3.font.size = Pt(12)
    run3.font.bold = True
    
    # Add symptoms section
    r7 = document.add_heading()
    run7 = r7.add_run("Patient's Symptoms")
    run7.font.name = "Century Gothic"
    run7.font.size = Pt(20)
    run7.font.color.rgb = RGBColor(79, 129, 189)
    run7.font.bold = True

    symptoms = [Symptom1, Symptom2, Symptom3, Symptom4, Symptom5]

    # Create a table with the required number of rows and columns
    # Start with a table that has just one row for the header
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'  # Optional: Set a table style

    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = "Symptom Number"
    header_cells[1].text = "Symptom Description"

    # Function to check and replace underscore
    def replace_underscore(symptom):
        symptom_text = symptom.get()  # Get the string value from the StringVar
        if not symptom_text.strip():  # If the symptom is empty, return None
            return None
        symptom_text = symptom_text.title()  # Capitalize the first letter of each word
        if "_" in symptom_text:  # Check if underscore is present
            return symptom_text.replace("_", " ")  # Replace underscore with space
        return symptom_text  # Return as is if no underscore

    # Set font for headers
    for cell in header_cells:
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = "Century Gothic"
        run.font.size = Pt(17)
        run.font.bold = True

    # Add symptom data to the table
    row_index = 1  # Start adding rows from index 1 (since index 0 is the header)
    for i, symptom in enumerate(symptoms, start=1):
        symptom_text = replace_underscore(symptom)  # Process the symptom text

        # Skip if symptom text is empty or None
        if symptom_text is None or not symptom_text.strip():
            continue

        # Add the symptom and its description to the table
        row_cells = table.add_row().cells  # Add a new row for the symptom
        row_cells[0].text = f"Symptom {i}"
        row_cells[1].text = symptom_text

        # Apply font styling to each cell
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Century Gothic"
                    run.font.size = Pt(10)
                    run.font.bold = True  # Apply bold to the description text
    
    r7 = document.add_heading()
    run7 = r7.add_run("Defination of Selected Symptoms") # New Heading
    run7.font.name = "Century Gothic"
    run7.font.size = Pt(20)
    run7.font.color.rgb = RGBColor(79, 129, 189)
    run7.font.bold = True
    
    def get_disease_info(symptoms):
        # Read the CSV file
        with open('Data\\diseases.csv', 'r') as csvfile:
            reader = csv.DictReader(csvfile)
            disease_info = {row['Disease']: row['Description'] for row in reader}

        # Fetch and return information for the given symptoms
        result = {}
        for symptom_var in symptoms:  # Assuming 'symptoms' is a list of StringVar objects
            symptom = symptom_var.get()  # Get the string value of the StringVar
            result[symptom] = disease_info.get(symptom, 'Information not available')

        return result


    symptoms = [Symptom1, Symptom2, Symptom3, Symptom4, Symptom5]

    info = get_disease_info(symptoms)
    table = document.add_table(rows=len(symptoms) + 1, cols=2)  

    # Style the table (optional)
    table.style = 'Table Grid'  

    # Add the header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Symptom"
    header_cells[1].text = "Information"    

    # Format the header
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Century Gothic"
                run.font.size = Pt(20)
                run.font.bold = True    

    # Add the symptoms and their associated information to the table
    row_index = 1  # Start from the second row (index 1) since the first row is the header
    for symptom in symptoms:
        symptom_text = symptom.get().strip()  # Get the string value from the StringVar and remove extra spaces
        # Skip empty symptoms
        if not symptom_text:
            continue
        
        disease_info = info.get(symptom_text, "Information not available")  # Fetch associated information
        # Add a new row for the symptom
        row_cells = table.rows[row_index].cells
        sys1 = symptom_text.replace("_", " ")
        row_cells[0].text = sys1.title()
        row_cells[1].text = disease_info
        # Format the cells
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Century Gothic"
                    run.font.size = Pt(10)
                    run.font.bold = True
    
        # Increment row index for the next valid symptom
        row_index += 1
    
    # Remove any unused rows at the end
    for _ in range(len(symptoms) + 1 - row_index):
        table._element.remove(table.rows[-1]._element)  # Remove extra rows
    
    r7 = document.add_heading()
    run7 = r7.add_run("Results")
    run7.font.name = "Century Gothic"
    run7.font.size = Pt(20)
    run7.font.color.rgb = RGBColor(255, 0, 0)
    run7.font.bold = True
   
    DecisionTree = canvas.itemcget(tid_1, "text")
    Randomforest = canvas.itemcget(tid_2, "text")
    NaiveBayes = canvas.itemcget(tid_3, "text")
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'  # Optional: Set a table style

    # Add the header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Algorithms"
    header_cells[1].text = "Disease Predicted"

    # Function to apply font styling
    def apply_font_styling(cell, font_name="Century Gothic", font_size=10, bold=True):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.font.bold = bold

    # Format the header row
    for cell in header_cells:
        apply_font_styling(cell, font_size=20, bold=True)

    # List of models and their descriptions
    models = [
        ("Decision Tree", DecisionTree),
        ("Random Forest", Randomforest),
        ("Naive Bayes", NaiveBayes)
    ]

    # Add data to the table, but only for models with non-empty descriptions
    for model, description in models:
        # Check if description is empty (or contains only whitespace)
        if description.strip() and description != ".":
            # Add a new row for the model with a non-empty description and does not contain "."
            row_cells = table.add_row().cells
            row_cells[0].text = model
            row_cells[1].text = description
            apply_font_styling(row_cells[0])  # Apply font styling to the model column
        apply_font_styling(row_cells[1])  # Apply font styling to the description column

    heading = document.add_heading('Recommendations for each Disease Predcited', level=1)
    def analyze_disease_and_generate_doc(disease_name):
        # Load the CSV file
        csv_file = "data\\diseases_recommended_actions.csv"
        try:
            with open(csv_file, mode='r') as file:
                reader = csv.DictReader(file)
                for row in reader:
                    if row['Disease'].lower() == disease_name.lower():
                        recommendation = row['Recommended Action']

                        # Add heading
                        heading_run = heading.runs[0]
                        heading_run.font.name = 'Century Gothic'
                        heading_run.font.size = Pt(20)
                        heading_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
    
                        # Add content
                        disease_paragraph = document.add_paragraph(f"Disease: {disease_name}")
                        recommendation_paragraph = document.add_paragraph(f"Recommended Action: {recommendation}")
    
                        for paragraph in [disease_paragraph, recommendation_paragraph]:
                            for run in paragraph.runs:
                                run.font.name = 'Century Gothic'
                                run.font.bold = True
                return f"Disease '{disease_name}' not found in the database."
        except FileNotFoundError:
            return "The CSV file is missing. Please ensure the file exists."
    models2 = [
        ("Decision Tree", DecisionTree),
        ("Random Forest", Randomforest),
        ("Naive Bayes", NaiveBayes)
    ]

    # Add data to the table, but only for models with non-empty descriptions
    for model, description in models2:
        # Check if description is empty (or contains only whitespace)
        if description.strip() and description != ".":
            # Add a new row for the model with a non-empty description and does not contain "."
            analyze_disease_and_generate_doc(description)
  
        
# Save the document
    folder_selected = filedialog.askdirectory(title="Select a location to save Report")
    document.save(f"{folder_selected}\\report.docx")
    messagebox.showinfo("Report Generated", "Report Generated Successfully")
    response = messagebox.askyesno("Report Generated", "Do you want to convert this report into pdf?")

 
    def convert_docx_to_pdf(docx_path, pdf_path=None):
        """
        Convert a DOCX file to PDF.

        :param docx_path: Path to the input DOCX file.
        :param pdf_path: Path to save the output PDF file. If None, saves in the same directory as the DOCX file.
        """
        try:
            convert(docx_path, pdf_path)
            print(f"Successfully converted '{docx_path}' to PDF.")
        except Exception as e:
            print(f"Error: {e}")
            
    if response:
        from reportgen import PDFViewer
        convert_docx_to_pdf(f"{folder_selected}\\report.docx", f"{folder_selected}\\report.pdf")
        messagebox.showinfo("Report Converted", "Report Converted Successfully")
        PDFViewer(window, f"{folder_selected}\\report.pdf")
    else:
        pass
    
    
button_image_5 = PhotoImage(
    file=relative_to_assets("button_5.png"))
button_5 = Button(
    image=button_image_5,
    borderwidth=0,
    highlightthickness=0,
    command=generate_report,
    relief="flat"
)






button_5.place(
    x=41.0,
    y=458.0,
    width=270.0,
    height=58.81910705566406
)
canvas.create_text(
    766.0,
    425.0,
    anchor="nw",
    text="RF  ",
    fill="#FFFFFF",
    font=("Lexend Bold", 27 * -1)
)

canvas.create_text(
    763.0,
    466.0,
    anchor="nw",
    text="NB",
    fill="#FFFFFF",
    font=("Lexend Bold", 27 * -1)
)

canvas.create_text(
    766.0,
    383.0,
    anchor="nw",
    text="DT ",
    fill="#FFFFFF",
    font=("Lexend Bold", 27 * -1)
)

canvas.create_rectangle(
    815.0,
    398.0,
    840.0,
    404.0,
    fill="#FFFFFF",
    outline="")

canvas.create_rectangle(
    815.0,
    480.0,
    840.0,
    486.0,
    fill="#FFFFFF",
    outline="")

canvas.create_rectangle(
    815.0,
    439.0,
    840.0,
    445.0,
    fill="#FFFFFF",
    outline="")

# Create comboboxes at the positions of the rectangles
create_combobox(496.0, 200.0, 696.0, 235.0 ,Symptom1, "TCombobox")
create_combobox(496.0, 262.0, 696.0, 297.0 ,Symptom2, "TCombobox")
create_combobox(496.0, 324.0, 696.0, 359.0 ,Symptom3, "TCombobox")
create_combobox(496.0, 448.0, 696.0, 483.0 ,Symptom4, "TCombobox")
create_combobox(496.0, 386.0, 696.0, 421.0 ,Symptom5, "TCombobox")


tid_1 = canvas.create_text(
    861.0,
    392.0,
    anchor="nw",
    text=".",
    fill="#FFFFFF",
    font=("Lexend Bold", 14 * -1)
)

tid_2 = canvas.create_text(
    861.0,
    433.0,
    anchor="nw",
    text=".",
    fill="#FFFFFF",
    font=("Lexend Bold", 14 * -1)
)

tid_3 = canvas.create_text(
    861.0,
    474.0,
    anchor="nw",
    text=".",
    fill="#FFFFFF",
    font=("Lexend Bold", 14 * -1)
)

canvas.create_text(
    54.0,
    203.0,
    anchor="nw",
    text="Name :",
    fill="#FFFFFF",
    font=("Lexend Regular", 20 * -1)
)

entry_image_2 = PhotoImage(
    file=relative_to_assets("entry_2.png"))
entry_bg_2 = canvas.create_image(
    222.5,
    275.5,
    image=entry_image_2
)
entry_2 = Text(
    
    bd=0,
    bg="#D9D9D9",
    fg="#000716",
    highlightthickness=0
)
entry_2.place(
    x=141.0,
    y=265.0,
    width=163.0,
    height=25
)

canvas.create_text(
    36.0,
    263.0,
    anchor="nw",
    text="Contact :",
    fill="#FFFFFF",
    font=("Lexend Regular", 20 * -1)
)
window.resizable(False, False)
window.mainloop()
